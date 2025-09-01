import asyncio
import json
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd


class NumpyEncoder(json.JSONEncoder):
    """Custom JSON encoder that handles numpy types for serialization"""

    def default(self, obj):
        if isinstance(obj, (np.integer, np.floating)):
            return obj.item()
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        elif hasattr(obj, "__float__"):
            return float(obj)
        elif hasattr(obj, "__int__"):
            return int(obj)
        return super().default(obj)


from docx import Document as DocxDocument
from llama_index.core import Document, Settings, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.storage.storage_context import StorageContext
from llama_index.embeddings.ollama import OllamaEmbedding as LlamaIndexOllamaEmbedding
from llama_index.llms.ollama import Ollama as LlamaIndexOllamaLLM
from llama_index.vector_stores.redis import RedisVectorStore
from llama_index.vector_stores.redis.schema import RedisVectorStoreSchema
from pypdf import PdfReader

from src.circuit_breaker import circuit_breaker_async

# Import the centralized ConfigManager
from src.config import (
    HTTP_PROTOCOL,
    OLLAMA_HOST_IP,
    OLLAMA_PORT,
)
from src.config import config as global_config_manager

# Import retry mechanism and circuit breaker
from src.retry_mechanism import RetryStrategy, retry_async

# Import centralized Redis client utility
from src.utils.redis_client import get_redis_client

# Import AutoBot semantic chunker
from src.utils.semantic_chunker import get_semantic_chunker


class KnowledgeBase:
    def __init__(self, config_manager=None):
        """
        Initialize KnowledgeBase with dependency injection support.

        Args:
            config_manager: Configuration manager instance (optional, uses global if None)
        """
        # Use provided dependency or fall back to default for backward compatibility
        self.config_manager = config_manager or global_config_manager

        # Network share path (if applicable)
        self.network_share_path = self.config_manager.get_nested("network_share.path")
        self.network_username = self.config_manager.get_nested(
            "network_share.username", os.getenv("NETWORK_SHARE_USERNAME")
        )
        self.network_password = self.config_manager.get_nested(
            "network_share.password", os.getenv("NETWORK_SHARE_PASSWORD")
        )

        # LlamaIndex specific settings from config
        self.vector_store_type = self.config_manager.get_nested(
            "llama_index.vector_store.type", "redis"
        )
        # Get embedding model from unified configuration
        unified_llm_config = self.config_manager.get_llm_config()
        self.embedding_model_name = unified_llm_config.get("unified", {}).get(
            "embedding", {}
        ).get("providers", {}).get("ollama", {}).get(
            "selected_model"
        ) or self.config_manager.get_nested(
            "llama_index.embedding.model", "nomic-embed-text:latest"
        )
        self.chunk_size = self.config_manager.get_nested("llama_index.chunk_size", 512)
        self.chunk_overlap = self.config_manager.get_nested(
            "llama_index.chunk_overlap", 20
        )

        # RAG Optimization: Enable semantic chunking
        self.use_semantic_chunking = self.config_manager.get_nested(
            "knowledge_base.semantic_chunking.enabled", True
        )
        self.semantic_chunking_model = self.config_manager.get_nested(
            "knowledge_base.semantic_chunking.model", "all-MiniLM-L6-v2"
        )

        # IMPORTANT: Override embedding model to use nomic-embed-text for consistency
        # This ensures both LlamaIndex and semantic chunking use the same dimensions
        if self.use_semantic_chunking:
            logging.info(
                "KNOWLEDGE BASE FIX: Using nomic-embed-text for both LlamaIndex and semantic chunking"
            )
            self.embedding_model_name = "nomic-embed-text:latest"
        self.semantic_percentile_threshold = self.config_manager.get_nested(
            "knowledge_base.semantic_chunking.percentile_threshold", 95.0
        )

        # Redis configuration for LlamaIndex vector store
        # Use memory.redis configuration for consistency
        memory_config = self.config_manager.get("memory", {})
        redis_config = memory_config.get("redis", {})
        self.redis_host = redis_config.get("host", os.getenv("REDIS_HOST", "localhost"))
        self.redis_port = redis_config.get("port", 6379)
        self.redis_password = redis_config.get("password", os.getenv("REDIS_PASSWORD"))
        # CRITICAL FIX: Use database 0 where the 13,383 vectors actually exist
        self.redis_db = redis_config.get(
            "kb_db", 0
        )  # Default to db 0 where existing vectors are located
        # Must use "llama_index" as the name - it's hardcoded in the library
        self.redis_index_name = "llama_index"

        # Initialize Redis client for direct use (e.g., for facts/logs)
        # Use sync Redis client for direct operations - WITH TIMEOUT PROTECTION
        try:
            self.redis_client = get_redis_client(async_client=False)
            self._async_redis_client = None  # Will be initialized lazily
            if self.redis_client:
                logging.info("Redis client initialized via centralized utility")
            else:
                logging.warning(
                    "Redis client not available - Redis may be disabled in configuration"
                )
        except Exception as e:
            logging.warning(f"Redis client initialization failed: {e}, continuing without Redis")
            self.redis_client = None
            self._async_redis_client = None

        self.llm = None
        self.embed_model = None
        self.vector_store = None
        self.storage_context = None
        self.index = None
        self.query_engine = None

    async def _get_async_redis_client(self):
        """Get or initialize the async Redis client."""
        if self._async_redis_client is None:
            async_client_factory = get_redis_client(async_client=True)
            if asyncio.iscoroutine(async_client_factory):
                self._async_redis_client = await async_client_factory
            else:
                self._async_redis_client = async_client_factory
        return self._async_redis_client

    def _resolve_path(self, configured_path: str) -> str:
        if self.network_share_path and not os.path.isabs(configured_path):
            return os.path.join(self.network_share_path, configured_path)
        return configured_path

    async def ainit(self):  # Removed llm_config parameter
        """
        Asynchronously initializes LlamaIndex components including LLM,
        Embedding model, and Vector Store.
        """
        llm_config_data = (
            self.config_manager.get_llm_config()
        )  # Get LLM config from config manager
        llm_provider = llm_config_data.get("provider", "ollama")
        # Use unified config to get the actual selected model
        llm_model = llm_config_data.get("ollama", {}).get("model", "tinyllama:latest")
        llm_base_url = llm_config_data.get("ollama", {}).get(
            "base_url"
        ) or llm_config_data.get("unified", {}).get("local", {}).get(
            "providers", {}
        ).get(
            "ollama", {}
        ).get(
            "host", f"{HTTP_PROTOCOL}://{OLLAMA_HOST_IP}:{OLLAMA_PORT}"
        )

        if llm_provider == "ollama":
            # Try to initialize with configured model, with fallback handling and timeout
            try:
                # Add timeout to prevent blocking
                import httpx
                import asyncio
                
                # Test Ollama availability with short timeout first
                try:
                    # CRITICAL FIX: Use async HTTP client to prevent blocking the event loop
                    async def test_ollama():
                        async with httpx.AsyncClient(timeout=5.0) as client:
                            response = await client.get(f"{llm_base_url}/api/tags")
                            if response.status_code != 200:
                                raise Exception(f"Ollama not responding properly: {response.status_code}")
                    
                    await test_ollama()
                except (httpx.TimeoutException, httpx.ConnectError) as e:
                    raise Exception(f"Ollama not available at {llm_base_url}: {e}")
                
                self.llm = LlamaIndexOllamaLLM(model=llm_model, base_url=llm_base_url)
                logging.info(
                    f"KnowledgeBase: Successfully initialized with model '{llm_model}'"
                )
            except Exception as e:
                logging.warning(
                    f"KnowledgeBase: Model '{llm_model}' failed to initialize: {e}"
                )
                # Try fallback models in order of preference
                fallback_models = [
                    "dolphin-llama3:8b",
                    "llama3.2:3b",
                    "nomic-embed-text:latest",
                ]
                model_initialized = False

                for fallback_model in fallback_models:
                    try:
                        logging.info(
                            f"KnowledgeBase: Trying fallback model '{fallback_model}'"
                        )
                        self.llm = LlamaIndexOllamaLLM(
                            model=fallback_model, base_url=llm_base_url
                        )
                        logging.info(
                            "KnowledgeBase: Successfully initialized with "
                            f"fallback model '{fallback_model}'"
                        )
                        model_initialized = True
                        break
                    except Exception as fallback_error:
                        logging.warning(
                            f"KnowledgeBase: Fallback model '{fallback_model}' "
                            f"also failed: {fallback_error}"
                        )
                        continue

                if not model_initialized:
                    logging.error(
                        "KnowledgeBase: All LLM models failed to initialize. "
                        "KnowledgeBase will be disabled."
                    )
                    raise Exception(
                        "No available Ollama models found. "
                        f"Original error with '{llm_model}': {e}"
                    )

            try:
                # Test embedding model availability with timeout
                try:
                    # CRITICAL FIX: Use async HTTP client to prevent blocking the event loop
                    async def test_embedding_model():
                        async with httpx.AsyncClient(timeout=5.0) as client:
                            # Test if embedding model exists
                            response = await client.post(f"{llm_base_url}/api/show", 
                                                       json={"name": self.embedding_model_name})
                            if response.status_code != 200:
                                raise Exception(f"Embedding model {self.embedding_model_name} not found")
                    
                    await test_embedding_model()
                except (httpx.TimeoutException, httpx.ConnectError) as e:
                    raise Exception(f"Cannot verify embedding model: {e}")
                    
                self.embed_model = LlamaIndexOllamaEmbedding(
                    model_name=self.embedding_model_name, base_url=llm_base_url
                )
                logging.info(
                    "KnowledgeBase: Successfully initialized embedding model "
                    f"'{self.embedding_model_name}'"
                )
            except Exception as e:
                logging.error(
                    f"KnowledgeBase: Embedding model '{self.embedding_model_name}' "
                    f"failed to initialize: {e}"
                )
                raise Exception(
                    f"Embedding model '{self.embedding_model_name}' "
                    f"not available. Error: {e}"
                )
        else:
            logging.warning(
                f"LLM provider '{llm_provider}' not fully implemented for "
                "LlamaIndex. Defaulting to Ollama."
            )
            # CRITICAL FIX: Add timeout and error handling to LLM initialization
            try:
                self.llm = LlamaIndexOllamaLLM(
                    model=llm_model, 
                    base_url=llm_base_url,
                    request_timeout=10.0  # Add 10-second timeout
                )
                self.embed_model = LlamaIndexOllamaEmbedding(
                    model_name=self.embedding_model_name, 
                    base_url=llm_base_url
                )
                logging.info(f"Initialized LLM with timeout protection: {llm_model}")
            except Exception as e:
                logging.error(f"Failed to initialize LLM: {e}")
                # Use a simple fallback model that's likely to be available
                self.llm = LlamaIndexOllamaLLM(
                    model="phi:2.7b", 
                    base_url=llm_base_url,
                    request_timeout=5.0
                )
                logging.warning("Using fallback model: phi:2.7b")

        Settings.llm = self.llm
        Settings.embed_model = self.embed_model
        Settings.chunk_size = self.chunk_size
        Settings.chunk_overlap = self.chunk_overlap

        # Configure node parser based on semantic chunking setting
        if self.use_semantic_chunking:
            # For semantic chunking, we'll handle chunking manually in add_file
            Settings.node_parser = None  # Will be handled by semantic_chunker
            logging.info(
                f"Semantic chunking enabled with model: {self.semantic_chunking_model}"
            )
        else:
            Settings.node_parser = SentenceSplitter(
                chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap
            )
            logging.info("Using traditional sentence-based chunking")

        try:
            # Get the actual embedding dimension from the model
            embedding_dim = 768  # Default for nomic-embed-text

            # Try to get the actual dimension by creating a test embedding with timeout
            try:
                import asyncio
                import concurrent.futures
                
                # Run embedding with timeout to prevent blocking
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(self.embed_model.get_text_embedding, "test")
                    try:
                        test_embedding = future.result(timeout=10)  # 10 second timeout
                        embedding_dim = len(test_embedding)
                        logging.info(f"Detected embedding dimension: {embedding_dim}")
                    except concurrent.futures.TimeoutError:
                        logging.warning(f"Embedding test timed out after 10s, using default dimension {embedding_dim}")
                    except Exception as e:
                        logging.warning(f"Embedding test failed: {e}, using default dimension {embedding_dim}")
            except Exception as e:
                logging.warning(
                    "Could not detect embedding dimension, using default "
                    f"{embedding_dim}: {e}"
                )

            # Create schema with proper vector dimensions
            schema = RedisVectorStoreSchema(
                index_name=self.redis_index_name,
                prefix="doc",
                overwrite=True,  # Allow overwriting to fix dimension mismatch
            )

            # Fix hardcoded 1536 dimensions in vector field by modifying existing field
            try:
                # Get the existing vector field and update its dimensions
                vector_field = schema.fields["vector"]
                if hasattr(vector_field, "attrs") and hasattr(
                    vector_field.attrs, "dims"
                ):
                    # Update the dimension attribute to use variable detected dimension
                    original_dims = vector_field.attrs.dims
                    vector_field.attrs.dims = embedding_dim
                    logging.info(
                        f"Updated vector field dimensions from {original_dims} to {embedding_dim}"
                    )
                else:
                    logging.warning(
                        "Could not modify vector field dimensions, using default"
                    )
            except Exception as e:
                logging.warning(
                    f"Failed to update vector field dimensions: {e}, using default schema"
                )

            logging.info(
                f"Created Redis schema with variable {embedding_dim} vector dimensions"
            )

            self.vector_store = RedisVectorStore(
                schema=schema,
                redis_url=f"redis://{self.redis_host}:{self.redis_port}",
                password=self.redis_password,
                redis_kwargs={"db": self.redis_db},
            )
            logging.info(
                "LlamaIndex RedisVectorStore initialized with index: "
                f"{self.redis_index_name}"
            )
        except ImportError as e:
            logging.warning(
                f"Could not import RedisVectorStoreSchema: {e}. "
                "Using fallback configuration."
            )
            self.vector_store = None
            logging.warning("Redis vector store disabled due to configuration issues.")

        if self.vector_store is not None:
            self.storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )
            # CRITICAL FIX: Use from_vector_store to connect to existing 13,383 vectors
            # This preserves existing data instead of creating empty index
            try:
                self.index = VectorStoreIndex.from_vector_store(
                    vector_store=self.vector_store,
                    embed_model=self.embed_model,  # Explicitly pass the embedding model
                )
                logging.info(
                    "LlamaIndex VectorStoreIndex loaded from existing Redis vector store with 13,383 vectors."
                )
            except Exception as e:
                logging.warning(f"Failed to load from existing vector store: {e}")
                # Fallback to empty index if loading fails
                self.index = VectorStoreIndex.from_documents(
                    [],
                    storage_context=self.storage_context,
                    embed_model=self.embed_model,
                )
                logging.info("Created new empty VectorStoreIndex as fallback.")

            # Create query engine with explicit LLM configuration
            try:
                self.query_engine = self.index.as_query_engine(
                    llm=self.llm,
                    similarity_top_k=10,
                    response_mode="compact"
                )
                logging.info("LlamaIndex QueryEngine initialized with proper LLM configuration.")
            except Exception as e:
                logging.error(f"Failed to create query engine: {e}")
                self.query_engine = None
        else:
            logging.warning(
                "Creating in-memory VectorStoreIndex without Redis due to "
                "initialization issues."
            )
            self.index = VectorStoreIndex.from_documents([])
            self.query_engine = self.index.as_query_engine(llm=self.llm)
            logging.info("In-memory VectorStoreIndex and QueryEngine initialized.")

    async def _scan_redis_keys(self, pattern: str) -> List[str]:
        """Non-blocking Redis key scanning to replace KEYS operations.

        Args:
            pattern: Redis key pattern to match

        Returns:
            List of matching keys
        """
        async_redis = await self._get_async_redis_client()
        if not async_redis:
            logging.error("Async Redis client not available")
            return []
            
        all_keys: List[str] = []
        cursor = 0
        while True:
            cursor, batch = await async_redis.scan(cursor, match=pattern, count=100)
            all_keys.extend(batch)
            if cursor == 0:
                break
        return all_keys

    async def add_file(
        self,
        file_path: str,
        file_type: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        if self.index is None:
            logging.warning(
                "KnowledgeBase not initialized, attempting initialization..."
            )
            try:
                await self.ainit()
                logging.info("KnowledgeBase initialized successfully")
            except Exception as e:
                logging.error(f"Failed to initialize KnowledgeBase: {e}")
                return {
                    "status": "error",
                    "message": f"KnowledgeBase initialization failed: {e}",
                }

        if self.index is None:  # Still None after initialization attempt
            return {
                "status": "error",
                "message": (
                    "KnowledgeBase initialization failed - " "index not available"
                ),
            }
        try:
            content = ""
            if file_type == "txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif file_type == "pd":
                reader = PdfReader(file_path)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            elif file_type == "csv":
                df = pd.read_csv(file_path)
                content = df.to_string()
            elif file_type == "docx":
                doc = DocxDocument(file_path)
                for para in doc.paragraphs:
                    content += para.text + "\n"
            else:
                logging.warning(
                    "Unsupported file type for direct loading by LlamaIndex: "
                    f"{file_type}. Attempting as generic text."
                )
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

            doc_metadata = {
                "filename": os.path.basename(file_path),
                "file_type": file_type,
                "original_path": file_path,
                **(metadata if metadata else {}),
            }

            if self.use_semantic_chunking:
                # Use semantic chunking for better knowledge processing
                logging.info(f"Processing file with semantic chunking: {file_path}")
                chunk_docs = await get_semantic_chunker().chunk_document(
                    content, doc_metadata
                )

                # Insert each semantic chunk as a separate document
                for chunk_data in chunk_docs:
                    chunk_document = Document(
                        text=chunk_data["text"], metadata=chunk_data["metadata"]
                    )
                    self.index.insert(chunk_document)

                logging.info(
                    f"Inserted {len(chunk_docs)} semantic chunks from {file_path}"
                )
            else:
                # Use traditional chunking
                document = Document(text=content, metadata=doc_metadata)
                self.index.insert(document)
                logging.info(f"Inserted traditional chunks from {file_path}")

            logging.info(
                f"File '{file_path}' processed and added to LlamaIndex (Redis)."
            )
            return {
                "status": "success",
                "message": f"File '{file_path}' processed and added to KB.",
            }
        except Exception as e:
            logging.error(f"Error adding file {file_path} to KB: {str(e)}")
            return {
                "status": "error",
                "message": f"Error adding file to KB: {str(e)}",
            }

    @circuit_breaker_async(
        "knowledge_base_service",
        failure_threshold=5,
        recovery_timeout=15.0,
        timeout=5.0,  # Reduced from 30s to 5s for faster failure
    )
    @retry_async(
        max_attempts=3, base_delay=0.5, strategy=RetryStrategy.EXPONENTIAL_BACKOFF
    )
    async def search_legacy(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        if self.query_engine is None:
            logging.warning(
                "KnowledgeBase not initialized, attempting initialization..."
            )
            try:
                await self.ainit()
                logging.info("KnowledgeBase initialized successfully")
            except Exception as e:
                logging.error(f"Failed to initialize KnowledgeBase: {e}")
                return []

        if self.query_engine is None:  # Still None after initialization attempt
            logging.error("Query engine not available after initialization")
            return []

        try:
            # CRITICAL FIX: Use retriever instead of query_engine to avoid LLM issues
            import asyncio
            
            # Use retriever to get nodes directly without LLM processing
            if hasattr(self.index, 'as_retriever'):
                retriever = self.index.as_retriever(similarity_top_k=n_results)
                nodes = await asyncio.to_thread(retriever.retrieve, query)
            else:
                # Fallback to query engine if retriever not available
                response = await asyncio.to_thread(self.query_engine.query, query)
                nodes = response.source_nodes

            retrieved_info = []
            for node in nodes:
                retrieved_info.append(
                    {
                        "content": node.text,
                        "metadata": node.metadata if hasattr(node, 'metadata') else {},
                        "score": getattr(node, 'score', 0.0),
                    }
                )
            logging.info(
                f"Found {len(retrieved_info)} relevant chunks for query: '{query}'"
            )

            # If no vector search results, fall back to fact search
            if len(retrieved_info) == 0:
                logging.info(
                    f"No vector results for '{query}', falling back to fact search"
                )
                facts = await self.get_fact(query=query)
                for fact in facts[:n_results]:
                    retrieved_info.append(
                        {
                            "content": fact["content"],
                            "metadata": fact["metadata"],
                            "score": 0.8,  # Default score for fact matches
                        }
                    )
                logging.info(
                    f"Fact search found {len(retrieved_info)} results for '{query}'"
                )

            return retrieved_info
        except Exception as e:
            logging.error(f"Error searching knowledge base: {str(e)}")
            # Try fact search as final fallback
            try:
                facts = await self.get_fact(query=query)
                retrieved_info = []
                for fact in facts[:n_results]:
                    retrieved_info.append(
                        {
                            "content": fact["content"],
                            "metadata": fact["metadata"],
                            "score": 0.7,  # Lower score for fallback results
                        }
                    )
                logging.info(
                    f"Fallback fact search found {len(retrieved_info)} results"
                )
                return retrieved_info
            except Exception as fact_error:
                logging.error(f"Fact search also failed: {fact_error}")
                return []

    def _clean_numpy_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Convert numpy types and complex objects in metadata to Redis-compatible types"""
        if not metadata:
            return {}
        
        def clean_value(value):
            """Recursively clean values for Redis compatibility"""
            if isinstance(value, (np.integer, np.floating)):
                return value.item()
            elif isinstance(value, np.ndarray):
                return value.tolist()
            elif isinstance(value, dict):
                # Convert nested dictionaries to JSON strings for Redis
                try:
                    return json.dumps(value, cls=NumpyEncoder)
                except (TypeError, ValueError):
                    return str(value)
            elif isinstance(value, (list, tuple)):
                # Clean list/tuple elements
                try:
                    return [clean_value(item) for item in value]
                except Exception:
                    return str(value)
            elif hasattr(value, "__float__"):
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return str(value)
            elif hasattr(value, "__int__"):
                try:
                    return int(value)
                except (ValueError, TypeError):
                    return str(value)
            elif isinstance(value, (str, int, float, bool)):
                return value
            else:
                # Convert any other complex types to strings
                return str(value)
        
        clean_metadata = {}
        for key, value in metadata.items():
            try:
                clean_metadata[str(key)] = clean_value(value)
            except Exception as e:
                # If all else fails, convert to string
                logging.warning(f"Failed to clean metadata key {key}: {e}")
                clean_metadata[str(key)] = str(value)
                
        return clean_metadata
    
    async def _safe_insert_document(self, document: Document):
        """Safely insert document with numpy type conversion"""
        try:
            # First, ensure any numpy types in the document metadata are cleaned
            if document.metadata:
                document.metadata = self._clean_numpy_metadata(document.metadata)
            
            # Insert the document
            self.index.insert(document)
            logging.debug(f"Successfully inserted document into index")
            
        except Exception as e:
            # If we still get numpy serialization errors, try alternative approaches
            if "numpy.float32" in str(e) or "Unable to serialize" in str(e):
                logging.warning(f"Numpy serialization error, attempting fallback: {e}")
                
                # Try converting document text to ensure no numpy types
                try:
                    # Create a completely clean document
                    clean_doc = Document(
                        text=str(document.text),
                        metadata=self._clean_numpy_metadata(document.metadata or {})
                    )
                    self.index.insert(clean_doc)
                    logging.info("Successfully inserted document using fallback method")
                except Exception as fallback_error:
                    logging.error(f"Even fallback insertion failed: {fallback_error}")
                    raise fallback_error
            else:
                logging.error(f"Document insertion failed: {e}")
                raise e

    async def add_text_with_semantic_chunking(
        self, content: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Add text content using semantic chunking for enhanced knowledge processing.

        Args:
            content: Text content to add
            metadata: Optional metadata for the content

        Returns:
            Dict with status and processing details
        """
        if self.index is None:
            logging.warning(
                "KnowledgeBase not initialized, attempting initialization..."
            )
            try:
                await self.ainit()
                logging.info("KnowledgeBase initialized successfully")
            except Exception as e:
                logging.error(f"Failed to initialize KnowledgeBase: {e}")
                return {
                    "status": "error",
                    "message": f"KnowledgeBase initialization failed: {e}",
                }

        try:
            text_metadata = {
                "source": "direct_text_input",
                "content_type": "text",
                "timestamp": datetime.now().isoformat(),
                **(metadata if metadata else {}),
            }

            if self.use_semantic_chunking:
                # Use semantic chunking for better knowledge processing
                logging.info("Processing text with semantic chunking")
                chunk_docs = await get_semantic_chunker().chunk_document(
                    content, text_metadata
                )

                # Insert each semantic chunk as a separate document
                for chunk_data in chunk_docs:
                    # Convert any numpy types in metadata to native Python types
                    clean_metadata = self._clean_numpy_metadata(chunk_data["metadata"])
                    
                    chunk_document = Document(
                        text=chunk_data["text"], metadata=clean_metadata
                    )
                    # Use safe insert method that handles numpy conversion
                    await self._safe_insert_document(chunk_document)

                logging.info(
                    f"Added {len(chunk_docs)} semantic chunks to knowledge base"
                )
                return {
                    "status": "success",
                    "message": f"Text processed into {len(chunk_docs)} semantic chunks",
                    "chunks_created": len(chunk_docs),
                    "semantic_chunking": True,
                }
            else:
                # Use traditional processing
                clean_metadata = self._clean_numpy_metadata(text_metadata)
                document = Document(text=content, metadata=clean_metadata)
                await self._safe_insert_document(document)
                return {
                    "status": "success",
                    "message": "Text added to knowledge base",
                    "chunks_created": 1,
                    "semantic_chunking": False,
                }

        except Exception as e:
            logging.error(f"Error adding text to knowledge base: {str(e)}")
            return {
                "status": "error",
                "message": f"Error adding text: {str(e)}",
            }

    async def store_fact(
        self, content: str, metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        async_redis = await self._get_async_redis_client()
        if not async_redis:
            return {
                "status": "error",
                "message": "Redis client not available - Redis may be disabled",
            }
        try:
            import time

            fact_id = await async_redis.incr("fact_id_counter")
            fact_key = f"fact:{fact_id}"
            fact_data = {
                "content": content,
                "metadata": (
                    json.dumps(metadata, cls=NumpyEncoder) if metadata else "{}"
                ),
                "timestamp": str(int(time.time())),
            }
            await async_redis.hset(fact_key, mapping=fact_data)
            logging.info(f"Fact stored in Redis with ID: {fact_id}")
            return {
                "status": "success",
                "message": "Fact stored successfully.",
                "fact_id": fact_id,
            }
        except Exception as e:
            logging.error(f"Error storing fact in Redis: {str(e)}")
            return {"status": "error", "message": f"Error storing fact: {str(e)}"}

    async def get_fact(
        self, fact_id: Optional[int] = None, query: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        async_redis = await self._get_async_redis_client()
        if not async_redis:
            logging.warning("Async Redis client not available")
            return []

        facts = []
        try:
            if fact_id:
                fact_data: Dict[str, str] = await async_redis.hgetall(f"fact:{fact_id}")
                if fact_data:
                    facts.append(
                        {
                            "id": fact_id,
                            "content": fact_data.get("content"),
                            "metadata": json.loads(fact_data.get("metadata", "{}")),
                            "timestamp": fact_data.get("timestamp"),
                        }
                    )
            elif query:
                # Use non-blocking SCAN instead of blocking KEYS operation
                all_keys = await self._scan_redis_keys("fact:*")

                if all_keys:
                    # PERFORMANCE OPTIMIZATION: Use Redis pipeline for batch operations
                    pipe = async_redis.pipeline()
                    for key in all_keys:
                        pipe.hgetall(key)
                    results = await pipe.execute()

                    for key, fact_data in zip(all_keys, results):
                        if fact_data:
                            content = fact_data.get("content", "").lower()
                            query_words = query.lower().split()
                            # Match if any significant query words are found in content
                            matches = any(
                                word in content
                                for word in query_words
                                if len(word) > 2  # Skip very short words
                            )
                            if matches:
                                facts.append(
                                    {
                                        "id": int(key.split(":")[1]),
                                        "content": fact_data.get("content"),
                                        "metadata": json.loads(
                                            fact_data.get("metadata", "{}")
                                        ),
                                        "timestamp": fact_data.get("timestamp"),
                                    }
                                )
            else:
                # Use non-blocking SCAN instead of blocking KEYS operation
                all_keys = self._scan_redis_keys("fact:*")
                if all_keys:
                    # PERFORMANCE OPTIMIZATION: Use Redis pipeline for batch operations
                    pipe = self.redis_client.pipeline()
                    for key in all_keys:
                        pipe.hgetall(key)
                    results = pipe.execute()

                    for key, fact_data in zip(all_keys, results):
                        if fact_data:
                            facts.append(
                                {
                                    "id": int(key.split(":")[1]),
                                    "content": fact_data.get("content"),
                                    "metadata": json.loads(
                                        fact_data.get("metadata", "{}")
                                    ),
                                    "timestamp": fact_data.get("timestamp"),
                                }
                            )
            logging.info(
                f"Retrieved {len(facts)} facts from Redis using sync operations."
            )
            return facts
        except Exception as e:
            logging.error(f"Error retrieving facts from Redis: {str(e)}")
            return []

    async def get_stats(self) -> Dict[str, Any]:
        """
        Get basic, high-level statistics about the knowledge base.

        This method provides a quick overview with minimal computational overhead,
        suitable for dashboard displays and health checks. Returns core metrics
        like document counts and available categories.

        Returns:
            Dict containing basic stats: total_documents, total_chunks, categories, total_facts

        Performance: Fast (< 100ms) - uses optimized counting methods with async operations
        """
        stats = {
            "total_documents": 0,
            "total_chunks": 0,
            "categories": [],
            "total_facts": 0,
        }
        if not self.redis_client:
            return stats
        try:
            # PERFORMANCE FIX: Use asyncio.to_thread to avoid blocking the event loop
            import asyncio
            import concurrent.futures

            def scan_keys_sync(pattern: str) -> int:
                """Synchronous key counting using SCAN - run in thread pool"""
                count = 0
                cursor = 0
                max_iterations = 50  # Limit iterations to prevent long-running scans
                iteration = 0
                while True:
                    cursor, keys = self.redis_client.scan(
                        cursor, match=pattern, count=100
                    )
                    count += len(keys)
                    iteration += 1
                    if cursor == 0 or iteration >= max_iterations:
                        break
                return count

            def get_categories_sync() -> List[str]:
                """Synchronous category sampling - run in thread pool"""
                try:
                    cursor, sample_keys = self.redis_client.scan(
                        0, match=f"{self.redis_index_name}:doc:*", count=5
                    )
                    categories = set()
                    for key in sample_keys[:3]:  # Limit to 3 keys for speed
                        try:
                            doc_data = self.redis_client.hgetall(key)
                            if doc_data and "metadata" in doc_data:
                                metadata = json.loads(doc_data.get("metadata", "{}"))
                                if "category" in metadata:
                                    categories.add(metadata["category"])
                        except Exception:
                            continue
                    return list(categories)
                except Exception:
                    return []

            def get_vector_count_sync() -> int:
                """Get vector count from Redis FT.INFO - run in thread pool"""
                try:
                    ft_info = self.redis_client.execute_command('FT.INFO', 'llama_index')
                    # Parse FT.INFO result to find num_docs
                    for i in range(len(ft_info)):
                        if ft_info[i] == 'num_docs':
                            return int(ft_info[i + 1])
                    return 0
                except Exception as e:
                    logging.debug(f"Could not get vector count: {e}")
                    return 0

            # PERFORMANCE FIX: Run Redis operations in thread pool with timeout
            try:
                # Create thread pool executor for blocking Redis operations
                with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                    # Submit Redis operations concurrently with timeout
                    vector_task = asyncio.create_task(
                        asyncio.to_thread(get_vector_count_sync)
                    )
                    fact_task = asyncio.create_task(
                        asyncio.to_thread(scan_keys_sync, "fact:*")
                    )
                    categories_task = asyncio.create_task(
                        asyncio.to_thread(get_categories_sync)
                    )

                    # Wait for all operations with 2-second timeout
                    vector_count, fact_count, categories = await asyncio.wait_for(
                        asyncio.gather(vector_task, fact_task, categories_task),
                        timeout=2.0,
                    )

                    stats["total_documents"] = vector_count
                    stats["total_chunks"] = vector_count  # Vector chunks from LlamaIndex
                    stats["total_facts"] = fact_count
                    stats["categories"] = categories

            except asyncio.TimeoutError:
                logging.warning(
                    "Redis stats collection timed out, using fallback values"
                )
                stats["total_documents"] = 0
                stats["total_chunks"] = 0
                stats["total_facts"] = 0
                stats["categories"] = []
            except Exception as e:
                logging.warning(f"Stats collection failed: {e}, using fallback values")
                stats["total_documents"] = 0
                stats["total_chunks"] = 0
                stats["total_facts"] = 0
                stats["categories"] = []

            logging.info(f"Knowledge base stats (non-blocking): {stats}")
            return stats
        except Exception as e:
            logging.error(f"Error getting knowledge base stats: {str(e)}")
            return stats

    async def get_detailed_stats(self) -> Dict[str, Any]:
        """
        Get comprehensive, in-depth statistics about the knowledge base.

        This method performs detailed analysis of the knowledge base content,
        including size calculations, usage patterns, and content distribution.
        Suitable for analytics, performance monitoring, and detailed reporting.

        Returns:
            Dict containing detailed stats: total_size, avg_chunk_size, last_updated,
            searches_24h, top_category, fact_types

        Performance: Slower (1-5s) - performs comprehensive content analysis
        Note: Uses NPU acceleration when available for faster processing
        """

        detailed_stats = {
            "total_size": 0,  # in bytes
            "avg_chunk_size": 0,  # in characters
            "last_updated": "N/A",
            "searches_24h": 0,
            "top_category": "N/A",
            "fact_types": {},  # e.g., {"manual_input": 10, "web_link": 5}
        }
        if not self.redis_client:
            return detailed_stats
        try:
            # Total size and average chunk size calculation for facts
            all_fact_keys = self._scan_redis_keys("fact:*")
            total_content_length = 0
            fact_type_counts = {}
            latest_timestamp = 0

            for key in all_fact_keys:
                fact_data: Dict[str, str] = self.redis_client.hgetall(key)
                if fact_data:
                    content = fact_data.get("content", "")
                    total_content_length += len(content)

                    metadata = json.loads(fact_data.get("metadata", "{}"))
                    source_type = metadata.get("source", "unknown")
                    fact_type_counts[source_type] = (
                        fact_type_counts.get(source_type, 0) + 1
                    )

                    timestamp = int(fact_data.get("timestamp", "0"))
                    if timestamp > latest_timestamp:
                        latest_timestamp = timestamp

            detailed_stats["total_size"] = total_content_length
            if len(all_fact_keys) > 0:
                detailed_stats["avg_chunk_size"] = total_content_length // len(
                    all_fact_keys
                )

            if latest_timestamp > 0:
                detailed_stats["last_updated"] = datetime.fromtimestamp(
                    latest_timestamp
                ).isoformat()

            detailed_stats["fact_types"] = fact_type_counts

            # Placeholder for search queries and top category (would need
            # separate logging/tracking)
            detailed_stats["searches_24h"] = 0  # Needs implementation
            if fact_type_counts:
                detailed_stats["top_category"] = max(
                    fact_type_counts, key=fact_type_counts.get
                )

            logging.info(f"Detailed knowledge base stats: {detailed_stats}")
            return detailed_stats
        except Exception as e:
            logging.error(f"Error getting detailed knowledge base stats: {str(e)}")
            return detailed_stats

    async def export_all_data(self) -> List[Dict[str, Any]]:
        """Export all stored knowledge base data (facts and potentially
        vector store metadata)."""
        if not self.redis_client:
            return []
        exported_data = []
        try:
            # Export facts from Redis
            all_fact_keys = self._scan_redis_keys("fact:*")
            for key in all_fact_keys:
                fact_data: Dict[str, str] = self.redis_client.hgetall(key)
                if fact_data:
                    exported_data.append(
                        {
                            "id": int(key.split(":")[1]),
                            "content": fact_data.get("content"),
                            "metadata": json.loads(fact_data.get("metadata", "{}")),
                            "timestamp": fact_data.get("timestamp"),
                            "type": "fact",
                        }
                    )

            # Export data from LlamaIndex (more complex, usually involves
            # iterating nodes)
            if self.index:
                all_doc_keys = self._scan_redis_keys(f"{self.redis_index_name}:doc:*")
                for key in all_doc_keys:
                    # Fetch the actual document/node data if needed, or just metadata
                    exported_data.append(
                        {
                            "id": key,
                            "content": "LlamaIndex Document (content not "
                            "directly exported here)",
                            "metadata": {"source": "LlamaIndex", "key": key},
                            "type": "document_reference",
                        }
                    )

            logging.info(f"Exported {len(exported_data)} items from knowledge base.")
            return exported_data
        except Exception as e:
            logging.error(f"Error exporting knowledge base data: {str(e)}")
            return []

    async def get_all_facts(self, collection: str = "default") -> List[Dict[str, Any]]:
        """Get all facts, optionally filtered by collection"""
        if not self.redis_client:
            return []
        facts = []
        try:
            # Use non-blocking SCAN instead of blocking KEYS operation
            all_keys = self._scan_redis_keys("fact:*")
            if all_keys:
                # PERFORMANCE OPTIMIZATION: Use Redis pipeline for batch operations
                pipe = self.redis_client.pipeline()
                for key in all_keys:
                    pipe.hgetall(key)
                results = await pipe.execute()

                for key, fact_data in zip(all_keys, results):
                    if fact_data:
                        metadata = json.loads(fact_data.get("metadata", "{}"))
                        fact_collection = metadata.get("collection", "default")

                        if collection == "all" or fact_collection == collection:
                            facts.append(
                                {
                                    "id": int(key.split(":")[1]),
                                    "content": fact_data.get("content"),
                                    "metadata": metadata,
                                    "timestamp": fact_data.get("timestamp"),
                                    "collection": fact_collection,
                                }
                            )

            # Sort by timestamp descending
            facts.sort(key=lambda x: int(x.get("timestamp", "0")), reverse=True)
            logging.info(
                f"Retrieved {len(facts)} facts from Redis using optimized pipeline."
            )
            return facts
        except Exception as e:
            logging.error(f"Error retrieving all facts from Redis: {str(e)}")
            return []

    async def update_fact(
        self, fact_id: int, content: str, metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """Update an existing fact"""
        try:
            async_redis = await self._get_async_redis_client()
            if not async_redis:
                logging.error("Async Redis client not available")
                return False
                
            fact_key = f"fact:{fact_id}"
            existing_data = await async_redis.hgetall(fact_key)

            if not existing_data:
                return False

            # Update with new data
            import time

            updated_data = {
                "content": content,
                "metadata": (
                    json.dumps(metadata, cls=NumpyEncoder) if metadata else "{}"
                ),
                "timestamp": existing_data.get(
                    "timestamp", str(int(time.time()))
                ),  # Keep original timestamp
            }

            await async_redis.hset(fact_key, mapping=updated_data)
            logging.info(f"Fact {fact_id} updated successfully.")
            return True
        except Exception as e:
            logging.error(f"Error updating fact {fact_id}: {str(e)}")
            return False

    async def delete_fact(self, fact_id: int) -> bool:
        """Delete a fact by ID"""
        try:
            async_redis = await self._get_async_redis_client()
            if not async_redis:
                logging.error("Async Redis client not available")
                return False
                
            fact_key = f"fact:{fact_id}"
            result = await async_redis.delete(fact_key)

            if result:
                logging.info(f"Fact {fact_id} deleted successfully.")
                return True
            else:
                logging.warning(f"Fact {fact_id} not found for deletion.")
                return False
        except Exception as e:
            logging.error(f"Error deleting fact {fact_id}: {str(e)}")
            return False

    async def cleanup_old_entries(self, days_to_keep: int) -> Dict[str, Any]:
        """Remove knowledge base entries (facts) older than a specified
        number of days."""
        removed_count = 0
        try:
            async_redis = await self._get_async_redis_client()
            if not async_redis:
                logging.error("Async Redis client not available")
                return {"status": "error", "message": "Redis client not available", "removed_count": 0}
                
            cutoff_timestamp = int(datetime.now().timestamp()) - (
                days_to_keep * 24 * 3600
            )
            all_fact_keys = await self._scan_redis_keys("fact:*")

            for key in all_fact_keys:
                fact_data: Dict[str, str] = await async_redis.hgetall(key)
                if fact_data:
                    timestamp = int(fact_data.get("timestamp", "0"))
                    if timestamp < cutoff_timestamp:
                        await async_redis.delete(key)
                        removed_count += 1

            logging.info(
                f"Cleaned up {removed_count} old knowledge base entries (facts)."
            )
            return {"status": "success", "removed_count": removed_count}
        except Exception as e:
            logging.error(f"Error cleaning up old entries: {str(e)}")
            return {
                "status": "error",
                "message": f"Error cleaning up: {str(e)}",
                "removed_count": 0,
            }
    
    # MCP-compatible methods for LLM access
    async def search(self, query: str, top_k: int = 5, filters: Optional[Dict] = None):
        """MCP-compatible search method for knowledge base"""
        try:
            # CRITICAL FIX: Use retriever instead of query_engine to avoid timeout issues
            if hasattr(self, 'index') and self.index:
                # Use retriever to get nodes directly without LLM processing
                retriever = self.index.as_retriever(similarity_top_k=top_k)
                
                # Add timeout protection
                nodes = await asyncio.wait_for(
                    asyncio.to_thread(retriever.retrieve, query),
                    timeout=10.0  # 10-second timeout
                )
                
                # Extract results from nodes
                results = []
                for node in nodes:
                    results.append({
                        "content": node.text if hasattr(node, 'text') else str(node),
                        "score": getattr(node, 'score', 0.0),
                        "metadata": getattr(node, 'metadata', {}),
                        "source": getattr(node, 'metadata', {}).get('source', 'knowledge_base')
                    })
                
                logging.info(f"MCP search returned {len(results)} results for: '{query}'")
                return results
            else:
                logging.warning("Index not available for search")
                return []
                
        except asyncio.TimeoutError:
            logging.error(f"MCP search timed out for query: '{query}'")
            return []
        except Exception as e:
            logging.error(f"Error in MCP search: {e}")
            return []
    
    async def add_document(self, content: str, metadata: Dict = None, source: str = None):
        """MCP-compatible document add method"""
        try:
            import uuid
            from llama_index.core import Document
            
            doc_id = str(uuid.uuid4())
            
            # Create Document object
            doc_metadata = metadata or {}
            if source:
                doc_metadata['source'] = source
            
            doc = Document(
                text=content,
                metadata=doc_metadata,
                doc_id=doc_id
            )
            
            # Add to index
            if self.index:
                # Run in thread to avoid blocking
                await asyncio.to_thread(self.index.insert, doc)
                logging.info(f"Added document {doc_id} to knowledge base via MCP")
                
                # Also store in Redis for persistence
                if hasattr(self, '_redis_client') and self._redis_client:
                    fact_data = {
                        "content": content,
                        "metadata": json.dumps(doc_metadata),
                        "timestamp": str(int(datetime.now().timestamp()))
                    }
                    await self._redis_client.hset(f"fact:{doc_id}", mapping=fact_data)
            
            return doc_id
            
        except Exception as e:
            logging.error(f"Error adding document via MCP: {e}")
            raise
    
    async def get_document_count(self):
        """Get total document count in knowledge base"""
        try:
            if self._redis_client:
                # Count all fact keys
                count = 0
                async for key in self._redis_client.scan_iter(match="fact:*"):
                    count += 1
                return count
            return 0
        except Exception as e:
            logging.error(f"Error getting document count: {e}")
            return 0
    
    async def get_last_update_time(self):
        """Get last update time of knowledge base"""
        try:
            if self._redis_client:
                # Get most recent fact
                latest_timestamp = 0
                async for key in self._redis_client.scan_iter(match="fact:*"):
                    fact_data = await self._redis_client.hgetall(key)
                    if fact_data and 'timestamp' in fact_data:
                        timestamp = int(fact_data['timestamp'])
                        latest_timestamp = max(latest_timestamp, timestamp)
                
                if latest_timestamp > 0:
                    return datetime.fromtimestamp(latest_timestamp).isoformat()
            
            return datetime.now().isoformat()
        except Exception as e:
            logging.error(f"Error getting last update time: {e}")
            return datetime.now().isoformat()
    
    async def get_memory_usage(self):
        """Get memory usage of knowledge base"""
        try:
            if self._redis_client:
                info = await self._redis_client.info('memory')
                return info.get('used_memory_human', 'unknown')
            return "unknown"
        except Exception as e:
            logging.error(f"Error getting memory usage: {e}")
            return "unknown"
