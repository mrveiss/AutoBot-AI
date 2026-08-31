# Copyright 2025-2026 mrveiss
# SPDX-License-Identifier: Apache-2.0
# AutoBot - AI-Powered Automation Platform
# Author: mrveiss
"""Table-capability reporting guards (#13895).

``tables: []`` used to mean two different things depending on input format, with
nothing in the result to tell them apart. PDF returned an empty list
unconditionally — the extractor was a stub whose comment said "would require
pdfplumber" — while DOCX walked ``doc.tables`` and did real work. A consumer
reading an identical-looking result could not distinguish *"this document has no
tables"* from *"we never looked"*, and the surrounding result reported
``confidence: 0.95`` either way.

``tables_attempted`` makes the distinction explicit. These tests hold it in
place, including on the failure paths, where nothing was extracted at all.
"""

import io

import pytest

from media.core.types import MediaInput, MediaType, ProcessingIntent
from media.document.extraction import ExtractedDocument, extract_docx
from media.document.pipeline import DocumentPipeline


def _pdf(pages: list) -> bytes:
    """Build a real PDF; ``None`` draws an image page, str draws text."""
    pytest.importorskip("reportlab", reason="reportlab needed to synthesize PDF fixtures")
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    for entry in pages:
        pdf.drawString(72, 720, entry or "")
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _docx(paragraphs: list, table_rows: list | None = None) -> bytes:
    """Build a real DOCX, optionally carrying one table."""
    pytest.importorskip("docx", reason="python-docx needed to synthesize DOCX fixtures")
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _result(raw: bytes, mime: str) -> dict:
    import asyncio

    media_input = MediaInput(
        media_id="doc",
        media_type=MediaType.DOCUMENT,
        intent=ProcessingIntent.ANALYSIS,
        data=raw,
        mime_type=mime,
        metadata={},
    )
    return asyncio.run(DocumentPipeline()._process_impl(media_input)).result_data


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# The reported defect
# ---------------------------------------------------------------------------


def test_pdf_now_looks_for_tables_and_says_so():
    """#14232 resolved what #13895 could only report.

    This test used to assert `tables_attempted is False` for PDFs, because
    `_extract_pdf_tables` was a stub returning `[]` while DOCX did real work —
    #13895 existed to make that emptiness honest rather than to fix it.

    PDF table extraction is implemented now, so the flag flips. The invariant
    #13895 established is unchanged and still asserted: an empty `tables` means
    "looked and found none" only when `tables_attempted` is true.
    """
    data = _result(_pdf(["some prose"]), "application/pdf")

    assert data["tables"] == []
    assert data["tables_attempted"] is True, "PDF table extraction runs now (#14232)"


def test_docx_reports_that_tables_were_attempted():
    data = _result(_docx(["intro"], [["a", "b"], ["c", "d"]]), _DOCX_MIME)

    assert data["tables_attempted"] is True
    assert data["tables"] == [[["a", "b"], ["c", "d"]]]


def test_a_docx_and_a_pdf_with_genuinely_no_tables_both_report_looking():
    """Both carry ``tables: []``, and since #14232 both looked.

    The flag no longer separates the two formats — it separates "looked and
    found none" from "did not look", which is what #13895 was for. A missing
    pdfplumber still produces ``attempted: False``; that case is covered in
    test_pdf_tables_14232.py.
    """
    docx_data = _result(_docx(["just prose"]), _DOCX_MIME)
    pdf_data = _result(_pdf(["just prose"]), "application/pdf")

    assert docx_data["tables"] == pdf_data["tables"] == []
    assert docx_data["tables_attempted"] is True, "DOCX looked and found none"
    assert pdf_data["tables_attempted"] is True, "PDF looks now too (#14232)"


# ---------------------------------------------------------------------------
# Failure paths — nothing was extracted, so nothing was attempted
# ---------------------------------------------------------------------------


def test_error_result_does_not_claim_tables_were_attempted():
    from unittest.mock import patch

    from media.document.extraction import DocumentExtractionError

    with patch("media.document.pipeline.extract_document", side_effect=DocumentExtractionError("bad")):
        data = _result(b"%PDF-1.4 garbage", "application/pdf")

    assert data["processing_status"] == "error"
    assert data["tables_attempted"] is False


def test_unavailable_result_does_not_claim_tables_were_attempted():
    from unittest.mock import patch

    from media.document.extraction import DocumentDependencyError

    with patch("media.document.pipeline.extract_document", side_effect=DocumentDependencyError("no lib")):
        data = _result(b"%PDF-1.4", "application/pdf")

    assert data["processing_status"] == "unavailable"
    assert data["tables_attempted"] is False


# ---------------------------------------------------------------------------
# Core dataclass
# ---------------------------------------------------------------------------


def test_extracted_document_defaults_to_not_attempted():
    """The safe default: a new format must opt in, not inherit a false claim."""
    assert ExtractedDocument(format="pdf", text="x").tables_attempted is False


def test_extract_docx_sets_the_flag_even_with_no_tables():
    """Attempted-and-found-none is the whole point of the flag."""
    extracted = extract_docx(_docx(["prose only"]))
    assert extracted.tables == ()
    assert extracted.tables_attempted is True


def test_plain_text_results_carry_no_table_fields_at_all():
    """Text has no table concept; inventing one would be a false negative."""
    data = _result(b"just text", "text/plain")
    assert "tables" not in data
    assert "tables_attempted" not in data
