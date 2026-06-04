"""Tests for DOCX (Microsoft Word) exporter."""

import io
import pytest
from docx import Document
from services.transcript_export.base import Segment, Transcript
from services.transcript_export.docx_exporter import DOCXExporter


@pytest.fixture
def sample_transcript():
    """Sample transcript for testing."""
    return Transcript(
        id="test-789",
        title="Board Meeting",
        duration_seconds=120.0,
        language="en",
        segments=[
            Segment(
                id="seg-1",
                transcript_id="test-789",
                start_time=0.0,
                end_time=5.0,
                speaker_label="CEO",
                text="Let's start the quarterly review.",
            ),
            Segment(
                id="seg-2",
                transcript_id="test-789",
                start_time=6.0,
                end_time=12.0,
                speaker_label="CFO",
                text="Our revenue increased by 15% this quarter.",
            ),
        ],
    )


@pytest.mark.asyncio
async def test_docx_generate(sample_transcript):
    """Test DOCX generation."""
    exporter = DOCXExporter(sample_transcript)
    result = await exporter.generate()

    # Parse generated DOCX
    doc = Document(io.BytesIO(result))

    # Check that document has content
    assert len(doc.paragraphs) > 0

    # Check title is in document
    text_content = "\n".join(p.text for p in doc.paragraphs)
    assert "Board Meeting" in text_content
    assert "CEO" in text_content
    assert "Let's start the quarterly review." in text_content
    assert "CFO" in text_content


@pytest.mark.asyncio
async def test_docx_mime_type():
    """Test DOCX MIME type."""
    transcript = Transcript(
        id="test",
        title="Test",
        duration_seconds=10.0,
        language="en",
        segments=[],
    )
    exporter = DOCXExporter(transcript)
    assert exporter.get_mime_type() == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.mark.asyncio
async def test_docx_file_extension():
    """Test DOCX file extension."""
    transcript = Transcript(
        id="test",
        title="Test",
        duration_seconds=10.0,
        language="en",
        segments=[],
    )
    exporter = DOCXExporter(transcript)
    assert exporter.get_file_extension() == ".docx"


@pytest.mark.asyncio
async def test_docx_metadata_header(sample_transcript):
    """Test that DOCX includes metadata header."""
    exporter = DOCXExporter(sample_transcript)
    result = await exporter.generate()

    doc = Document(io.BytesIO(result))
    text_content = "\n".join(p.text for p in doc.paragraphs)

    # Check for metadata
    assert "Duration:" in text_content
    assert "Language:" in text_content
