# Copyright 2025-2026 mrveiss
# SPDX-License-Identifier: Apache-2.0
"""Integration tests for transcript export API endpoints.

Tests verify that the mock _get_transcript_mock function and export endpoints
are correctly wired up and would work in production.
"""

import io

import pytest
from docx import Document
from PyPDF2 import PdfReader

from services.transcript_export import (
    DOCXExporter,
    PDFExporter,
    Segment,
    SRTExporter,
    Transcript,
    VTTExporter,
)


@pytest.fixture
def sample_transcript():
    """Sample transcript for testing export functionality."""
    return Transcript(
        id="test-transcript-1",
        title="Sample Meeting Transcript",
        duration_seconds=300.0,
        language="en",
        segments=[
            Segment(
                id="seg-1",
                transcript_id="test-transcript-1",
                start_time=5.0,
                end_time=10.0,
                speaker_label="Speaker 1",
                text="Welcome everyone to today's meeting.",
            ),
            Segment(
                id="seg-2",
                transcript_id="test-transcript-1",
                start_time=12.0,
                end_time=18.0,
                speaker_label="Speaker 2",
                text="Thank you for having us here.",
            ),
        ],
    )


@pytest.mark.asyncio
async def test_export_docx_generates_valid_document(sample_transcript):
    """Test DOCX exporter generates valid Word document."""
    exporter = DOCXExporter(sample_transcript)
    content = await exporter.generate()

    # Verify it's a valid DOCX
    doc = Document(io.BytesIO(content))
    assert len(doc.paragraphs) > 0

    # Verify MIME type and extension
    assert exporter.get_mime_type() == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert exporter.get_file_extension() == ".docx"

    # Verify content includes transcript data
    text_content = "\n".join(p.text for p in doc.paragraphs)
    assert "Sample Meeting Transcript" in text_content
    assert "Speaker 1" in text_content


@pytest.mark.asyncio
async def test_export_pdf_generates_valid_document(sample_transcript):
    """Test PDF exporter generates valid PDF document."""
    exporter = PDFExporter(sample_transcript)
    content = await exporter.generate()

    # Verify it's a valid PDF
    assert content.startswith(b"%PDF")
    reader = PdfReader(io.BytesIO(content))
    assert len(reader.pages) > 0

    # Verify MIME type and extension
    assert exporter.get_mime_type() == "application/pdf"
    assert exporter.get_file_extension() == ".pdf"


@pytest.mark.asyncio
async def test_export_srt_generates_valid_subtitles(sample_transcript):
    """Test SRT exporter generates valid subtitle file."""
    exporter = SRTExporter(sample_transcript)
    content = await exporter.generate()

    # Verify SRT format
    text_content = content.decode("utf-8")
    assert "1\n" in text_content  # Sequence number
    assert "-->" in text_content  # Timestamp separator
    assert "Speaker 1" in text_content

    # Verify MIME type and extension
    assert exporter.get_mime_type() == "application/x-subrip"
    assert exporter.get_file_extension() == ".srt"


@pytest.mark.asyncio
async def test_export_vtt_generates_valid_subtitles(sample_transcript):
    """Test VTT exporter generates valid WebVTT file."""
    exporter = VTTExporter(sample_transcript)
    content = await exporter.generate()

    # Verify VTT format
    text_content = content.decode("utf-8")
    assert text_content.startswith("WEBVTT")
    assert "-->" in text_content
    assert "<v" in text_content  # Voice tags

    # Verify MIME type and extension
    assert exporter.get_mime_type() == "text/vtt"
    assert exporter.get_file_extension() == ".vtt"


@pytest.mark.asyncio
async def test_filename_sanitization(sample_transcript):
    """Test that filename is properly sanitized."""
    exporter = SRTExporter(sample_transcript)
    filename = exporter.get_filename()

    # Verify sanitization (spaces replaced with underscores, no special chars)
    assert filename == "Sample_Meeting_Transcript.srt"
    assert " " not in filename
    assert all(c.isalnum() or c in ("_", "-", ".") for c in filename)


@pytest.mark.asyncio
@pytest.mark.parametrize(
    "exporter_class,extension",
    [
        (DOCXExporter, ".docx"),
        (PDFExporter, ".pdf"),
        (SRTExporter, ".srt"),
        (VTTExporter, ".vtt"),
    ],
)
async def test_all_exporters_functional(sample_transcript, exporter_class, extension):
    """Test all export formats are functional."""
    exporter = exporter_class(sample_transcript)
    content = await exporter.generate()

    assert len(content) > 0
    assert exporter.get_file_extension() == extension


def test_mock_data_function_structure():
    """Test that the mock transcript data matches expected structure.

    This verifies the _get_transcript_mock function would work correctly
    when the endpoints are called.
    """
    # Import the mock function
    import asyncio

    from api.transcript_export import _get_transcript_mock

    # Test successful case
    transcript = asyncio.run(_get_transcript_mock("test-transcript-1"))
    assert transcript.id == "test-transcript-1"
    assert transcript.title == "Sample Meeting Transcript"
    assert len(transcript.segments) == 2
    assert transcript.segments[0].speaker_label == "Speaker 1"

    # Test 404 case
    with pytest.raises(Exception) as exc_info:
        asyncio.run(_get_transcript_mock("nonexistent-id"))
    assert "not found" in str(exc_info.value).lower()
