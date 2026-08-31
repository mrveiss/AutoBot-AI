# Copyright 2025-2026 mrveiss
# SPDX-License-Identifier: Apache-2.0
# AutoBot - AI-Powered Automation Platform
# Author: mrveiss
"""Table-ingest guards (#14970).

PDF table extraction ran on every PDF (`extraction.py::_pdf_tables`, #14232) and
reached no ingest path: `content_extraction.py` and `document_parser.py` both
returned `extracted.text` and dropped `extracted.tables` on the floor for PDF,
while DOCX forked its own local join. `render_tables`/`render_text_and_tables`
(`media/document/provenance.py`) are the one renderer every table-bearing
consumer now shares.

DOCX fixtures run locally (`python-docx` is present). PDF fixtures need `pypdf`
and `pdfplumber`, neither installed in this sandbox — those tests are written
against the same real-fixture pattern the merged #13893/#14232 suites use and
are marked `importorskip`, so they run in CI, not here.
"""

import io

import pytest

from autobot_shared.ssot_config import config
from knowledge.connectors.content_extraction import extract_text_from_docx, extract_text_from_pdf
from media.document.extraction import ExtractedDocument, extract_docx
from media.document.provenance import (
    DEFAULT_MAX_TABLE_CHARS,
    TABLE_SECTION_MARKER,
    max_table_chars,
    render_tables,
    render_text_and_tables,
)
from utils.document_extractors import DocumentExtractor
from utils.document_parser import DocumentParser


def _docx_bytes_with_table() -> bytes:
    pytest.importorskip("docx", reason="python-docx needed to synthesize a DOCX fixture")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Widget"
    table.cell(1, 1).text = "5"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _pdf_bytes_with_table() -> bytes:
    """A PDF whose table has ruling lines, which is what pdfplumber's default
    layout-detection strategy needs to see a table at all (#14232) — an
    unstyled ``Table()`` renders as bare positioned text, indistinguishable
    from ordinary prose to pdfplumber's line/rectangle analysis.
    """
    pytest.importorskip("reportlab", reason="reportlab needed to synthesize a PDF fixture")
    pytest.importorskip("pypdf", reason="pypdf needed to read the PDF back")
    pytest.importorskip("pdfplumber", reason="pdfplumber needed to detect the table")
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    table = Table([["Name", "Amount"], ["Widget", "5"]])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    doc.build([table])
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# render_tables / render_text_and_tables — pure, no optional deps
# ---------------------------------------------------------------------------


def test_render_tables_pipe_joins_cells_and_blank_joins_tables():
    tables = (
        [["Name", "Amount"], ["Widget", "5"]],
        [["Other"]],
    )
    rendered = render_tables(tables)
    assert rendered == "Name | Amount\nWidget | 5\n\nOther"


def test_render_tables_drops_rows_with_no_non_blank_cell():
    tables = ([["Name", "Amount"], ["", ""], ["Widget", "5"]],)
    rendered = render_tables(tables)
    assert "\n\n" not in rendered.replace("Widget | 5", "")  # no blank-row artifact
    assert rendered == "Name | Amount\nWidget | 5"


def test_render_tables_returns_empty_string_for_no_tables():
    assert render_tables(()) == ""


def test_render_text_and_tables_appends_marker_when_both_present():
    document = ExtractedDocument(format="docx", text="intro", tables=([["a", "b"]],), tables_attempted=True)
    rendered = render_text_and_tables(document)
    assert rendered == f"intro\n\n{TABLE_SECTION_MARKER}\na | b"


def test_render_text_and_tables_passes_text_through_when_no_tables():
    document = ExtractedDocument(format="docx", text="just text", tables=(), tables_attempted=True)
    assert render_text_and_tables(document) == "just text"


def test_render_text_and_tables_handles_table_only_document():
    """A table-only PDF/DOCX has empty ``text`` — the marker must not lead a blank line."""
    document = ExtractedDocument(format="pdf", text="", tables=([["a", "b"]],), tables_attempted=True)
    rendered = render_text_and_tables(document)
    assert rendered == f"{TABLE_SECTION_MARKER}\na | b"
    assert not rendered.startswith("\n")


# ---------------------------------------------------------------------------
# max_table_chars — configurable, not a literal
# ---------------------------------------------------------------------------


def test_max_table_chars_is_configurable_not_a_literal(monkeypatch):
    monkeypatch.setattr(config.misc, "document_max_table_chars", "", raising=False)
    assert max_table_chars() == DEFAULT_MAX_TABLE_CHARS

    monkeypatch.setattr(config.misc, "document_max_table_chars", "500", raising=False)
    assert max_table_chars() == 500


@pytest.mark.parametrize("bad", ["not-a-number", "-1", "0"])
def test_invalid_max_table_chars_falls_back_to_the_default(monkeypatch, bad):
    monkeypatch.setattr(config.misc, "document_max_table_chars", bad, raising=False)
    assert max_table_chars() == DEFAULT_MAX_TABLE_CHARS


def test_render_tables_is_bounded_by_max_table_chars(monkeypatch):
    monkeypatch.setattr(config.misc, "document_max_table_chars", "10", raising=False)
    tables = ([["a very long cell value that exceeds the configured bound"]],)
    rendered = render_tables(tables)
    assert len(rendered) == 10


# ---------------------------------------------------------------------------
# DOCX: every consumer folds tables in, identically (#14970 AC)
# ---------------------------------------------------------------------------


def test_docx_table_reaches_the_connector_ingest_path():
    text = extract_text_from_docx(_docx_bytes_with_table())
    assert TABLE_SECTION_MARKER in text
    assert "Widget | 5" in text


def test_docx_table_reaches_the_document_parser_ingest_path(tmp_path):
    path = tmp_path / "report.docx"
    path.write_bytes(_docx_bytes_with_table())

    parser_text = DocumentParser()._parse_docx(path, {})
    assert TABLE_SECTION_MARKER in parser_text
    assert "Widget | 5" in parser_text


def test_docx_table_text_is_identical_via_connector_and_parser_paths(tmp_path):
    """#14970 AC: the same table produces the same text down every ingest path."""
    raw = _docx_bytes_with_table()
    path = tmp_path / "report.docx"
    path.write_bytes(raw)

    assert extract_text_from_docx(raw) == DocumentParser()._parse_docx(path, {})


@pytest.mark.asyncio
async def test_docx_table_reaches_the_batch_ingest_path(tmp_path):
    path = tmp_path / "report.docx"
    path.write_bytes(_docx_bytes_with_table())

    text = await DocumentExtractor.extract_from_docx(path)
    assert TABLE_SECTION_MARKER in text
    assert "Widget" in text and "Amount" in text


def test_docx_table_only_document_still_has_usable_content():
    """A table-only DOCX was already a successful extraction (#13884) — still is."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Solo cell"
    buffer = io.BytesIO()
    doc.save(buffer)

    extracted = extract_docx(buffer.getvalue())
    assert extracted.has_usable_content
    assert "Solo cell" in render_text_and_tables(extracted)


# ---------------------------------------------------------------------------
# PDF: needs pypdf + pdfplumber, absent from this sandbox — CI-only
# ---------------------------------------------------------------------------


def test_pdf_with_only_a_table_ingests_with_retrievable_content():
    """#14970 AC: a PDF containing only a table ingests with retrievable content."""
    text = extract_text_from_pdf(_pdf_bytes_with_table())
    assert TABLE_SECTION_MARKER in text
    assert "Widget" in text and "Amount" in text


def test_pdf_table_text_is_identical_via_connector_and_parser_paths(tmp_path):
    """#14970 AC: the same table produces the same text down every ingest path."""
    raw = _pdf_bytes_with_table()
    path = tmp_path / "report.pdf"
    path.write_bytes(raw)

    assert extract_text_from_pdf(raw) == DocumentParser()._parse_pdf(path, {})
