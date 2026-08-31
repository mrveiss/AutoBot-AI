# Copyright 2025-2026 mrveiss
# SPDX-License-Identifier: Apache-2.0
# AutoBot - AI-Powered Automation Platform
# Author: mrveiss
"""Canonical document text extraction (#13893).

Five independent PDF extractors existed before this module: the knowledge-base
upload endpoint, the media document pipeline, the Drive/OneDrive connectors, and
two fully orphaned helpers. Three were live, they disagreed on page markers, and
one still ran the unmaintained ``PyPDF2`` while the repo pins ``pypdf`` as a
security update. Every fix to the document path had to ship five times.

This module is the single implementation. Callers adapt its result to their own
error shape — an ``HTTPException`` for the API, a ``ProcessingResult`` for the
media pipeline, an empty string for the connectors — but nobody re-implements the
extraction itself.

Page text is kept **structured** rather than pre-joined. Callers that need a flat
string call :func:`render_pages`; callers that need per-page provenance read
``pages`` directly. That split is what lets page numbers reach retrieval as
metadata instead of as marker text baked into the embedding (#13894).

Page-offset lookups (``page_for_offset``/``pages_for_span``/``chunk_page_map``)
and table rendering (``render_tables``/``render_text_and_tables``) live in
:mod:`media.document.provenance` (#14970) — split out once folding tables into
ingest text pushed this module over ``MAX_LINES``. ``PageSpan`` and
``render_plain`` stayed here rather than moving with them, since
``api/knowledge.py`` imports ``render_plain`` directly and this way that file
needs no import-path change. Nothing in this module calls into
``provenance.py``, so the split carries no circular import.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from typing import Any, Dict, List, Mapping, Sequence, Tuple

from autobot_shared.env_utils import blank_to_none
from autobot_shared.logging_manager import get_logger
from autobot_shared.ssot_config import config

logger = get_logger(__name__)

# The one page-marker convention. Three different ones existed across the forks
# (``## Page N``, ``--- Page N ---``, and none at all on the live upload path).
PAGE_MARKER_TEMPLATE = "## Page {number}"

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK"
_DOCX_MARKER = b"word/"
_DOCX_SNIFF_BYTES = 2000

# #13884: fraction of pages that must carry text before an extraction counts as
# usable. Default 0.5 — a document where most pages are unreadable is a scan,
# whatever the remaining pages contain. Override when a corpus is legitimately
# mixed (title pages, plates, appendices of figures).
DEFAULT_MIN_TEXT_PAGE_RATIO = 0.5

# #13884 finding 1: the ratio above counts a page as readable when it carries a
# single character, which a scanner/DMS/Bates page-number stamp satisfies on
# every page. Measured against synthesized fixtures (reportlab + PIL): a
# "Page N of 10" stamp averages ~13 characters/page, a Bates+"CONFIDENTIAL"
# stamp ~25; a genuine single-field born-digital page (an invoice with five
# short lines) averages ~109, and ordinary dense prose ~3900. 50 sits between
# the stamp cluster and the real-content cluster with margin on both sides.
DEFAULT_MIN_CHARS_PER_PAGE = 50.0


def min_text_page_ratio() -> float:
    """Resolve the usable-text-layer ratio threshold from config.

    Resolved per call rather than captured at import so a deployment can retune
    it without a restart, and so tests can exercise the bounds.
    """
    raw = blank_to_none(config.misc.document_min_text_page_ratio)
    if raw is None:
        return DEFAULT_MIN_TEXT_PAGE_RATIO
    try:
        value = float(raw)
    except ValueError:
        logger.warning(
            "AUTOBOT_DOCUMENT_MIN_TEXT_PAGE_RATIO=%r is not a number; falling back to %s",
            raw,
            DEFAULT_MIN_TEXT_PAGE_RATIO,
        )
        return DEFAULT_MIN_TEXT_PAGE_RATIO
    if not 0.0 <= value <= 1.0:
        logger.warning(
            "AUTOBOT_DOCUMENT_MIN_TEXT_PAGE_RATIO=%s must be within [0.0, 1.0]; falling back to %s",
            value,
            DEFAULT_MIN_TEXT_PAGE_RATIO,
        )
        return DEFAULT_MIN_TEXT_PAGE_RATIO
    return value


def min_chars_per_page() -> float:
    """Resolve the characters-per-page floor from config.

    Companion to :func:`min_text_page_ratio`, same ``blank_to_none`` handling
    and the same "fall back to the default with a warning" contract for a
    misconfigured value (#13884).
    """
    raw = blank_to_none(config.misc.document_min_chars_per_page)
    if raw is None:
        return DEFAULT_MIN_CHARS_PER_PAGE
    try:
        value = float(raw)
    except ValueError:
        logger.warning(
            "AUTOBOT_DOCUMENT_MIN_CHARS_PER_PAGE=%r is not a number; falling back to %s",
            raw,
            DEFAULT_MIN_CHARS_PER_PAGE,
        )
        return DEFAULT_MIN_CHARS_PER_PAGE
    if value < 0:
        logger.warning(
            "AUTOBOT_DOCUMENT_MIN_CHARS_PER_PAGE=%s must be >= 0; falling back to %s",
            value,
            DEFAULT_MIN_CHARS_PER_PAGE,
        )
        return DEFAULT_MIN_CHARS_PER_PAGE
    return value


class DocumentExtractionError(Exception):
    """Raised when a document cannot be parsed.

    Callers map this to their own failure shape. It deliberately does not carry
    an HTTP status or a pipeline result — the core has no opinion about how a
    caller reports failure.
    """


class DocumentDependencyError(DocumentExtractionError):
    """Raised when the parsing library for a format is not installed.

    Kept distinct from a parse failure because the two need different responses:
    a corrupt file is the user's problem, a missing library is ours. The media
    pipeline reports this as ``unavailable`` rather than ``error`` so a
    deployment gap cannot be mistaken for a bad upload.
    """


@dataclass(frozen=True)
class PageText:
    """Text recovered from a single page, 1-indexed to match human page numbers."""

    number: int
    text: str


@dataclass(frozen=True)
class PageSpan:
    """Where one page's text sits in a rendered string, as ``[start, end)``.

    This is the provenance carrier: page numbers travel as offsets alongside the
    text rather than as markers inside it, so retrieval can cite a page without
    the citation having polluted the embedding (#13894).
    """

    number: int
    start: int
    end: int

    def as_metadata(self) -> Dict[str, int]:
        """Flat form for storage alongside a fact."""
        return {"page": self.number, "start": self.start, "end": self.end}


@dataclass(frozen=True)
class ExtractedDocument:
    """Structured result of a document extraction."""

    format: str
    text: str
    pages: Tuple[PageText, ...] = ()
    page_count: int | None = None
    tables: Tuple[Any, ...] = ()
    info: Mapping[str, str] = field(default_factory=dict)
    # #13895: whether table extraction was actually run for this format. Without
    # it an empty ``tables`` means both "this document has no tables" and "we
    # never looked", and the caller cannot tell which — PDF always returned []
    # while DOCX did real work, from an identical-looking result.
    tables_attempted: bool = False

    @property
    def char_count(self) -> int:
        """Number of characters of recovered text."""
        return len(self.text)

    @property
    def has_text(self) -> bool:
        """Whether any non-whitespace text was recovered.

        A PDF whose pages carry no text layer extracts to ``""`` rather than
        failing, so emptiness is a normal result that callers must check — it is
        the signal a scanned document needs OCR (#13884, #13896).
        """
        return bool(self.text.strip())

    @property
    def empty_page_numbers(self) -> Tuple[int, ...]:
        """Pages that carried no recoverable text, 1-indexed.

        Non-empty only for paginated formats. This is what lets a caller say
        *which* pages need OCR instead of failing or succeeding wholesale.
        """
        return tuple(page.number for page in self.pages if not page.text.strip())

    @property
    def text_page_ratio(self) -> float:
        """Fraction of pages that carried text; ``1.0`` for unpaginated formats.

        A 40-page scan with one stray text page reports ``has_text`` as True,
        which is technically correct and practically misleading — this is the
        measure that separates the two.
        """
        if not self.pages:
            return 1.0 if self.has_text else 0.0
        return (len(self.pages) - len(self.empty_page_numbers)) / len(self.pages)

    @property
    def avg_chars_per_page(self) -> float:
        """Average characters of raw page text, page markers excluded.

        :attr:`text_page_ratio` only asks whether a page has *any* text — one
        character qualifies, which a page-number stamp, Bates number, or
        filename footer satisfies on every page of a scan. This measures *how
        much* text landed per page, which is what actually separates a
        stamped scan from a page of real content (#13884).
        """
        if not self.pages:
            return float(self.char_count)
        return sum(len(page.text) for page in self.pages) / len(self.pages)

    @property
    def has_usable_text_layer(self) -> bool:
        """Whether enough of the document was readable to treat it as extracted.

        Distinct from :attr:`has_text`: a document can carry text and still be
        unusable, which is the case a scanned PDF with an OCR cover page hits,
        or one stamped with a page number or Bates number on every page.
        Unpaginated formats (no ``pages``) have no ratio or per-page floor to
        apply, so any recovered text is usable there.
        """
        if not self.has_text:
            return False
        if not self.pages:
            return True
        return self.text_page_ratio >= min_text_page_ratio() and self.avg_chars_per_page >= min_chars_per_page()

    @property
    def has_usable_content(self) -> bool:
        """Whether the extraction recovered anything a consumer can use.

        Distinct from :attr:`has_usable_text_layer`: a DOCX whose content is
        entirely a table has no text layer at all and is still a complete,
        successful extraction — its data lives in ``tables`` (#13884). PDF
        table extraction is implemented as of #14232; :attr:`tables_attempted`
        still reports whether it ran, rather than leaving callers to infer it
        from an empty list (#13895).
        """
        return self.has_usable_text_layer or bool(self.tables)


def render_pages(pages: Sequence[PageText], marker: str = PAGE_MARKER_TEMPLATE) -> str:
    """Join pages into one string, prefixing each with the canonical marker.

    Pages with no recovered text are skipped so a scanned page does not emit a
    bare marker with nothing under it.
    """
    parts = []
    for page in pages:
        if page.text.strip():
            parts.append(f"{marker.format(number=page.number)}\n{page.text}")
    return "\n\n".join(parts)


PAGE_SEPARATOR = "\n\n"


def render_plain(pages: Sequence[PageText], separator: str = PAGE_SEPARATOR) -> Tuple[str, Tuple[PageSpan, ...]]:
    """Join pages with **no** markers, returning the text and where each page sits.

    The marker-carrying :func:`render_pages` is the wrong input for an embedding:
    ``## Page 7`` is structure, not meaning, and it competes with the document's
    own words for similarity. This is the same text with the provenance moved out
    of the string and into character spans a caller stores as metadata (#13894).

    Empty pages are skipped, exactly as :func:`render_pages` skips them, so the
    two renderings stay page-for-page comparable.
    """
    parts: List[str] = []
    spans: List[PageSpan] = []
    offset = 0
    for page in pages:
        if not page.text.strip():
            continue
        if parts:
            offset += len(separator)
        spans.append(PageSpan(number=page.number, start=offset, end=offset + len(page.text)))
        parts.append(page.text)
        offset += len(page.text)
    return separator.join(parts), tuple(spans)


def detect_format(raw: bytes, mime_type: str = "") -> str:
    """Detect document format from magic bytes, falling back to MIME type."""
    mime = (mime_type or "").lower()
    if raw[:4] == _PDF_MAGIC:
        return "pdf"
    if raw[:2] == _ZIP_MAGIC and _DOCX_MARKER in raw[:_DOCX_SNIFF_BYTES]:
        return "docx"
    if "pdf" in mime:
        return "pdf"
    if "docx" in mime or "officedocument.wordprocessingml" in mime:
        return "docx"
    return "text"


def extract_pdf(raw: bytes) -> ExtractedDocument:
    """Extract the text layer from a PDF.

    A page whose text layer is absent contributes an empty :class:`PageText`
    rather than being dropped, so ``page_count`` stays truthful and a caller can
    tell *which* pages were unreadable.
    """
    reader = _pdf_reader(raw)
    pages = tuple(PageText(number=n, text=_pdf_page_text(page, n)) for n, page in enumerate(reader.pages, start=1))
    info = reader.metadata or {}
    tables, tables_attempted = _pdf_tables(raw)
    return ExtractedDocument(
        format="pdf",
        text=render_pages(pages),
        pages=pages,
        page_count=len(pages),
        info=_pdf_info(info),
        tables=tables,
        tables_attempted=tables_attempted,
    )


DEFAULT_MAX_TABLE_PAGES = 50


def max_table_pages() -> int:
    """Resolve how many pages table detection may scan (#14232).

    `page.extract_tables()` runs pdfplumber's line and rectangle layout
    analysis, which is materially heavier than pypdf's text extraction — and
    unlike OCR, which only runs on pages that already failed to produce text,
    this runs on every page of every PDF. A large table-dense document well
    inside the upload size limit can still carry hundreds of pages.

    So it is bounded the same way OCR's page ceiling is, and from the same kind
    of env-backed knob rather than a literal at the call site.
    """
    raw = blank_to_none(config.misc.document_max_table_pages)
    if raw is None:
        return DEFAULT_MAX_TABLE_PAGES
    try:
        value = int(raw)
    except ValueError:
        logger.warning(
            "AUTOBOT_DOCUMENT_MAX_TABLE_PAGES=%r is not an integer; using %d",
            raw,
            DEFAULT_MAX_TABLE_PAGES,
        )
        return DEFAULT_MAX_TABLE_PAGES
    if value <= 0:
        logger.warning(
            "AUTOBOT_DOCUMENT_MAX_TABLE_PAGES=%d is not positive; using %d",
            value,
            DEFAULT_MAX_TABLE_PAGES,
        )
        return DEFAULT_MAX_TABLE_PAGES
    return value


def _pdf_tables(raw: bytes) -> Tuple[Tuple[Any, ...], bool]:
    """Extract tables with pdfplumber, reporting whether it ran (#14232).

    Returns ``(tables, attempted)``. ``attempted`` is the field that keeps an
    empty result honest: ``([], True)`` means the document has no tables, and
    ``([], False)`` means nothing looked. #13895 introduced that distinction
    because PDF returned a bare ``[]`` unconditionally while DOCX did real work.

    pdfplumber is guard-imported. It is a heavier dependency than pypdf and a
    deployment may not carry it; a missing library degrades to
    ``attempted: False`` rather than failing an extraction whose text layer is
    perfectly readable.
    """
    try:
        import pdfplumber
    except ImportError:
        logger.debug("pdfplumber is not installed; PDF table extraction skipped")
        return (), False

    max_pages = max_table_pages()
    try:
        with pdfplumber.open(io.BytesIO(raw)) as document:
            tables = [
                _normalize_table(table)
                for page in document.pages[:max_pages]
                for table in (page.extract_tables() or [])
                if table
            ]
    except Exception as exc:
        # A text layer that read fine must not be lost because table detection
        # tripped, so this reports "did not look" rather than raising.
        logger.warning("PDF table extraction failed: %s", exc)
        return (), False

    return tuple(tables), True


def _normalize_table(table: Sequence[Sequence[Any]]) -> List[List[str]]:
    """Render one table as rows of cell strings.

    Identical in shape to :func:`_docx_table`, so a consumer needs one parser
    for both formats rather than branching on where the table came from.
    pdfplumber yields ``None`` for an empty cell where python-docx yields ``""``.
    """
    return [[("" if cell is None else str(cell)).strip() for cell in row] for row in table]


def _pdf_reader(raw: bytes) -> Any:
    """Build a pypdf reader, translating library failures to our own error."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentDependencyError("PDF support requires the pypdf library") from exc

    try:
        return PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise DocumentExtractionError(f"Failed to parse PDF: {exc}") from exc


def _pdf_page_text(page: Any, number: int) -> str:
    """Extract one page, degrading to empty text rather than failing the document."""
    try:
        return page.extract_text() or ""
    except Exception as exc:
        logger.warning("Failed to extract text from PDF page %s: %s", number, exc)
        return ""


def _pdf_info(info: Mapping[str, Any]) -> Dict[str, str]:
    """Normalize pypdf's slash-prefixed metadata keys."""
    return {
        "title": str(info.get("/Title", "") or ""),
        "author": str(info.get("/Author", "") or ""),
        "subject": str(info.get("/Subject", "") or ""),
        "creator": str(info.get("/Creator", "") or ""),
    }


def extract_docx(raw: bytes) -> ExtractedDocument:
    """Extract paragraphs, tables and core properties from a DOCX."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentDependencyError("DOCX support requires the python-docx library") from exc

    try:
        doc = DocxDocument(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = tuple(_docx_table(table) for table in doc.tables)
    except Exception as exc:
        raise DocumentExtractionError(f"Failed to parse DOCX: {exc}") from exc

    return ExtractedDocument(
        format="docx",
        text="\n".join(paragraphs),
        tables=tables,
        info=_docx_info(doc),
        tables_attempted=True,
    )


def _docx_table(table: Any) -> List[List[str]]:
    """Render one DOCX table as rows of cell strings."""
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _docx_info(doc: Any) -> Dict[str, str]:
    """Read DOCX core properties, tolerating documents that carry none."""
    try:
        props = doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
        }
    except Exception:
        return {}


def extract_plain_text(raw: bytes) -> ExtractedDocument:
    """Decode a plain-text document as UTF-8, falling back to latin-1."""
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    return ExtractedDocument(format="text", text=text)


def extract_document(raw: bytes, mime_type: str = "") -> ExtractedDocument:
    """Detect the format and extract, dispatching to the format-specific reader."""
    detected = detect_format(raw, mime_type)
    if detected == "pdf":
        return extract_pdf(raw)
    if detected == "docx":
        return extract_docx(raw)
    return extract_plain_text(raw)
