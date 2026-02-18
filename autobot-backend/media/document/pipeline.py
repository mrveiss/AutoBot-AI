# AutoBot - AI-Powered Automation Platform
# Copyright (c) 2025 mrveiss
# Author: mrveiss
#
# Document Processing Pipeline
# Issue #735: Organize media processing into dedicated pipelines
# Issue #932: Implement actual document processing

"""Document processing pipeline for text documents, PDFs, DOCX, etc."""

import base64
import io
import logging
from typing import Any, Dict, List

from media.core.pipeline import BasePipeline
from media.core.types import MediaInput, MediaType, ProcessingResult

# PDF support via pypdf (required)
try:
    from pypdf import PdfReader

    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

# DOCX support via python-docx (required)
try:
    from docx import Document as DocxDocument

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentPipeline(BasePipeline):
    """Pipeline for processing document content (PDF, DOCX, TXT, etc.)."""

    def __init__(self):
        """Initialize document processing pipeline."""
        super().__init__(
            pipeline_name="document",
            supported_types=[MediaType.DOCUMENT, MediaType.TEXT],
        )

    async def _process_impl(self, media_input: MediaInput) -> ProcessingResult:
        """Process document content."""
        result_data = await self._process_document(media_input)
        confidence = self._calculate_confidence(result_data)

        return ProcessingResult(
            result_id=f"document_{media_input.media_id}",
            media_id=media_input.media_id,
            media_type=media_input.media_type,
            intent=media_input.intent,
            success=True,
            confidence=confidence,
            result_data=result_data,
            processing_time=0.0,  # Set by BasePipeline
        )

    async def _process_document(self, media_input: MediaInput) -> Dict[str, Any]:
        """Dispatch to format-specific processor based on MIME type or content."""
        raw_bytes = self._decode_input(media_input.data)
        mime_type = (media_input.mime_type or "").lower()
        detected = self._detect_format(raw_bytes, mime_type)

        if detected == "pdf":
            return self._extract_pdf(raw_bytes, media_input.metadata)
        if detected == "docx":
            return self._extract_docx(raw_bytes, media_input.metadata)
        return self._extract_text(raw_bytes, mime_type, media_input.metadata)

    # ------------------------------------------------------------------
    # Decoding helpers
    # ------------------------------------------------------------------

    def _decode_input(self, data: Any) -> bytes:
        """Normalize input to bytes (base64 string, bytes, or file path)."""
        if isinstance(data, bytes):
            return data
        if isinstance(data, str):
            # Try base64 first, then treat as file path
            try:
                return base64.b64decode(data)
            except Exception:
                with open(data, "rb") as fh:
                    return fh.read()
        raise ValueError(f"Unsupported document data type: {type(data)}")

    def _detect_format(self, raw_bytes: bytes, mime_type: str) -> str:
        """Detect document format from magic bytes or MIME type."""
        if raw_bytes[:4] == b"%PDF":
            return "pdf"
        if raw_bytes[:2] == b"PK" and b"word/" in raw_bytes[:2000]:
            return "docx"
        if "pdf" in mime_type:
            return "pdf"
        if "docx" in mime_type or "officedocument.wordprocessingml" in mime_type:
            return "docx"
        return "text"

    # ------------------------------------------------------------------
    # Format-specific extractors
    # ------------------------------------------------------------------

    def _extract_pdf(self, raw_bytes: bytes, metadata: Dict) -> Dict[str, Any]:
        """Extract text and metadata from a PDF file using pypdf."""
        if not _PYPDF_AVAILABLE:
            return self._unavailable_result("pdf", "pypdf not installed", metadata)

        try:
            reader = PdfReader(io.BytesIO(raw_bytes))
            pages_text = [page.extract_text() or "" for page in reader.pages]
            full_text = "\n\n".join(pages_text)
            info = reader.metadata or {}
            tables = self._extract_pdf_tables(reader)
            return {
                "type": "document_analysis",
                "format": "pdf",
                "extracted_text": full_text,
                "page_count": len(reader.pages),
                "tables": tables,
                "document_info": {
                    "title": info.get("/Title", ""),
                    "author": info.get("/Author", ""),
                    "subject": info.get("/Subject", ""),
                    "creator": info.get("/Creator", ""),
                },
                "confidence": 0.95,
                "metadata": metadata,
            }
        except Exception as exc:
            logger.warning("PDF extraction failed: %s", exc)
            return self._error_result("pdf", str(exc), metadata)

    def _extract_pdf_tables(self, reader: Any) -> List[Dict]:
        """Best-effort table extraction from PDF annotations."""
        # pypdf doesn't natively extract tables; return placeholder structure
        # Full table extraction would require pdfplumber
        return []

    def _extract_docx(self, raw_bytes: bytes, metadata: Dict) -> Dict[str, Any]:
        """Extract text and tables from a DOCX file using python-docx."""
        if not _DOCX_AVAILABLE:
            return self._unavailable_result(
                "docx", "python-docx not installed", metadata
            )

        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            tables = self._extract_docx_tables(doc)
            core_props = self._extract_docx_properties(doc)
            return {
                "type": "document_analysis",
                "format": "docx",
                "extracted_text": full_text,
                "page_count": None,  # python-docx doesn't expose page count
                "paragraph_count": len(paragraphs),
                "tables": tables,
                "document_info": core_props,
                "confidence": 0.95,
                "metadata": metadata,
            }
        except Exception as exc:
            logger.warning("DOCX extraction failed: %s", exc)
            return self._error_result("docx", str(exc), metadata)

    def _extract_docx_tables(self, doc: Any) -> List[List[List[str]]]:
        """Extract tables from DOCX as list-of-rows-of-cells."""
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
        return tables

    def _extract_docx_properties(self, doc: Any) -> Dict[str, str]:
        """Extract core document properties (author, title, etc.)."""
        try:
            props = doc.core_properties
            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
            }
        except Exception:
            return {}

    def _extract_text(
        self, raw_bytes: bytes, mime_type: str, metadata: Dict
    ) -> Dict[str, Any]:
        """Extract text from plain text file (UTF-8 with latin-1 fallback)."""
        try:
            text = raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = raw_bytes.decode("latin-1", errors="replace")

        lines = text.splitlines()
        return {
            "type": "document_analysis",
            "format": "text",
            "extracted_text": text,
            "page_count": None,
            "line_count": len(lines),
            "char_count": len(text),
            "confidence": 1.0,
            "metadata": metadata,
        }

    # ------------------------------------------------------------------
    # Error/fallback helpers
    # ------------------------------------------------------------------

    def _unavailable_result(
        self, fmt: str, reason: str, metadata: Dict
    ) -> Dict[str, Any]:
        """Return structured result when a required library is unavailable."""
        logger.warning("Document pipeline (%s): %s", fmt, reason)
        return {
            "type": "document_analysis",
            "format": fmt,
            "extracted_text": "",
            "page_count": 0,
            "tables": [],
            "processing_status": "unavailable",
            "unavailability_reason": reason,
            "confidence": 0.0,
            "metadata": metadata,
        }

    def _error_result(self, fmt: str, error: str, metadata: Dict) -> Dict[str, Any]:
        """Return structured result on processing error."""
        return {
            "type": "document_analysis",
            "format": fmt,
            "extracted_text": "",
            "page_count": 0,
            "tables": [],
            "processing_status": "error",
            "error": error,
            "confidence": 0.0,
            "metadata": metadata,
        }

    def _calculate_confidence(self, result_data: Dict[str, Any]) -> float:
        """Calculate confidence score from result data."""
        return result_data.get("confidence", 0.5)
