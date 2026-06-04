# autobot-backend/transcriber/export/docx_export.py
# AutoBot - AI-Powered Automation Platform
# Copyright (c) 2025 mrveiss
# Author: mrveiss
"""Export transcript to Word (.docx) format using python-docx."""

import io

from docx import Document
from docx.shared import Pt, RGBColor

_SPEAKER_COLORS = [
    RGBColor(0x1A, 0x73, 0xE8),
    RGBColor(0xD9, 0x34, 0x25),
    RGBColor(0x18, 0x8A, 0x38),
    RGBColor(0xFB, 0xBC, 0x04),
    RGBColor(0x8A, 0x2B, 0xE2),
]


def _fmt_ts(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def build_docx(
    title: str,
    segments: list[dict],
    *,
    include_timestamps: bool,
    include_notes: bool,
    include_speaker_names: bool,
) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    speaker_index: dict[str, int] = {}
    for seg in segments:
        speaker = seg.get("speaker_name", "Unknown")
        if speaker not in speaker_index:
            speaker_index[speaker] = len(speaker_index)
        color = _SPEAKER_COLORS[speaker_index[speaker] % len(_SPEAKER_COLORS)]
        p = doc.add_paragraph()
        if include_speaker_names:
            run = p.add_run(f"{speaker}")
            run.bold = True
            run.font.color.rgb = color
            p.add_run("  ")
        if include_timestamps:
            p.add_run(f"[{_fmt_ts(seg['start'])} → {_fmt_ts(seg['end'])}]  ").italic = True
        p.add_run(seg["text"])
        if include_notes and seg.get("notes"):
            for note in seg["notes"]:
                np = doc.add_paragraph(style="Quote")
                np.add_run(f"📝 {note['content']}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
