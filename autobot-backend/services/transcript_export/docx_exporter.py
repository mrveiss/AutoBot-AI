"""DOCX (Microsoft Word) format exporter.

Generates a formatted Word document with:
- Title and metadata header
- Speaker labels in bold
- Timestamps in monospace
- Proper paragraph spacing
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.transcript_export.base import BaseExporter, Transcript


class DOCXExporter(BaseExporter):
    """Export transcript to DOCX (Microsoft Word) format."""

    def _format_time(self, seconds: float) -> str:
        """Format time as [MM:SS].

        Args:
            seconds: Time in seconds

        Returns:
            str: Formatted time string
        """
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"[{minutes:02d}:{secs:02d}]"

    def _format_duration(self, seconds: float) -> str:
        """Format duration as MM:SS or HH:MM:SS.

        Args:
            seconds: Duration in seconds

        Returns:
            str: Formatted duration string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        return f"{minutes:02d}:{secs:02d}"

    async def generate(self) -> bytes:
        """Generate DOCX file content.

        Returns:
            bytes: DOCX file content as bytes
        """
        doc = Document()

        # Add title
        title = doc.add_heading(self.transcript.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        metadata = doc.add_paragraph()
        metadata.add_run("Duration: ").bold = True
        metadata.add_run(self._format_duration(self.transcript.duration_seconds))
        metadata.add_run("\nLanguage: ").bold = True
        metadata.add_run(self.transcript.language.upper())
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add separator
        doc.add_paragraph("_" * 60)

        # Add segments
        for segment in self.transcript.segments:
            p = doc.add_paragraph()

            # Add timestamp
            timestamp_run = p.add_run(self._format_time(segment.start_time) + " ")
            timestamp_run.font.name = "Courier New"
            timestamp_run.font.size = Pt(10)
            timestamp_run.font.color.rgb = RGBColor(100, 100, 100)

            # Add speaker label
            speaker_run = p.add_run(f"{segment.speaker_label}: ")
            speaker_run.bold = True
            speaker_run.font.size = Pt(11)

            # Add text
            text_run = p.add_run(segment.text)
            text_run.font.size = Pt(11)

            # Add notes if present
            if segment.notes:
                notes_p = doc.add_paragraph()
                notes_run = notes_p.add_run(f"    Note: {segment.notes}")
                notes_run.italic = True
                notes_run.font.size = Pt(10)
                notes_run.font.color.rgb = RGBColor(80, 80, 80)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def get_mime_type(self) -> str:
        """Return DOCX MIME type."""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def get_file_extension(self) -> str:
        """Return DOCX file extension."""
        return ".docx"
